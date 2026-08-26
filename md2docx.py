import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown(md_path, base_dir):
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # 去掉 frontmatter
    if text.startswith('---'):
        end = text.find('---', 3)
        if end != -1:
            text = text[end + 3:].lstrip('\n')

    lines = text.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            blocks.append(('heading', level, m.group(2).strip()))
            i += 1
            continue

        # 图片
        m = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
        if m and line.strip().startswith('!'):
            alt = m.group(1)
            src = m.group(2)
            # Hexo 路径 /images/xxx -> source/images/xxx
            if src.startswith('/images/'):
                src = src.lstrip('/')
                src = os.path.join(base_dir, src)
            blocks.append(('image', alt, src))
            i += 1
            continue

        # 无序列表
        if re.match(r'^[-*+]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[-*+]\s+', lines[i]):
                item_text = re.sub(r'^[-*+]\s+', '', lines[i]).strip()
                # 去掉行首的数字加.（有序列表伪装）
                item_text = re.sub(r'^\d+\.\s*', '', item_text)
                items.append(item_text)
                i += 1
            blocks.append(('ul', items))
            continue

        # 有序列表
        if re.match(r'^\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i]):
                item_text = re.sub(r'^\d+\.\s+', '', lines[i]).strip()
                items.append(item_text)
                i += 1
            blocks.append(('ol', items))
            continue

        # 普通段落（合并连续行直到空行或块级元素）
        para_lines = []
        while i < len(lines):
            l = lines[i]
            if not l.strip():
                break
            if re.match(r'^(#{1,6})\s+', l):
                break
            if l.strip().startswith('!') and re.match(r'!\[.*?\]\(.*?\)', l.strip()):
                break
            if re.match(r'^[-*+]\s+', l) or re.match(r'^\d+\.\s+', l):
                break
            para_lines.append(l.rstrip())
            i += 1
        para_text = ' '.join(para_lines).strip()
        if para_text:
            blocks.append(('paragraph', para_text))

    return blocks


def add_inline_text(paragraph, text):
    # 处理 **加粗** 和普通文本
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def blocks_to_docx(blocks, output_path, base_dir):
    doc = Document()

    # 默认中文字体设置
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)

    for block in blocks:
        btype = block[0]

        if btype == 'heading':
            level = block[1]
            content = block[2]
            if level == 1:
                p = doc.add_heading(content, level=1)
            elif level == 2:
                p = doc.add_heading(content, level=2)
            else:
                p = doc.add_heading(content, level=min(level, 4))

        elif btype == 'paragraph':
            p = doc.add_paragraph()
            add_inline_text(p, block[1])

        elif btype == 'image':
            alt, src = block[1], block[2]
            if os.path.exists(src):
                try:
                    doc.add_picture(src, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if alt:
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = cap.add_run(alt)
                        r.italic = True
                except Exception as e:
                    doc.add_paragraph(f'[图片加载失败: {alt} - {e}]')
            else:
                doc.add_paragraph(f'[图片缺失: {alt} ({src})]')

        elif btype == 'ul':
            for item in block[1]:
                p = doc.add_paragraph(style='List Bullet')
                add_inline_text(p, item)

        elif btype == 'ol':
            for item in block[1]:
                p = doc.add_paragraph(style='List Number')
                add_inline_text(p, item)

    doc.save(output_path)


def main():
    if len(sys.argv) < 2:
        print('用法: python md2docx.py <md文件路径> [输出docx路径]')
        sys.exit(1)

    md_path = sys.argv[1]
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        output_path = os.path.splitext(md_path)[0] + '.docx'

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(md_path)))  # source 目录
    if not os.path.isdir(os.path.join(base_dir, 'images')):
        # 如果找不到 images 就用当前目录
        base_dir = os.path.dirname(os.path.abspath(md_path))

    blocks = parse_markdown(md_path, base_dir)
    blocks_to_docx(blocks, output_path, base_dir)
    print(f'已生成: {output_path}')


if __name__ == '__main__':
    main()
